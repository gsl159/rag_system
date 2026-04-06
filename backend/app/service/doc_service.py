"""
文档处理服务（轻量版）
替换 unstructured + spacy → pymupdf + python-docx + 标准库
依赖链：极简，无 spacy/numba/lxml 等重依赖

流程：上传 → MinIO → 解析 → 清洗 → 分块 → 质量评估 → 向量入库
"""
import re
import uuid
from pathlib import Path
from typing import List, Dict, Any

from sqlalchemy import update
from sqlalchemy.ext.asyncio import AsyncSession

from app.config.settings import settings
from app.utils.logger import logger
from app.core.generator import embed_client
from app.repository.postgres import Document, Chunk
from app.repository.vector_store import milvus_db


# ── 文档解析器 ────────────────────────────────

class DocParser:
    """
    轻量文档解析器
    PDF   → pymupdf (fitz)
    DOCX  → python-docx
    HTML  → 标准库 html.parser
    TXT/MD → 直接读取
    """

    def parse(self, file_path: str) -> str:
        suffix = Path(file_path).suffix.lower()
        try:
            if suffix == ".pdf":
                return self._parse_pdf(file_path)
            elif suffix in (".docx", ".doc"):
                return self._parse_docx(file_path)
            elif suffix in (".html", ".htm"):
                return self._parse_html(file_path)
            else:
                # TXT / MD / 其他文本格式
                return self._parse_text(file_path)
        except Exception as e:
            logger.error(f"解析失败 [{suffix}] {file_path}: {e}")
            return ""

    def _parse_pdf(self, path: str) -> str:
        """pymupdf 解析 PDF，提取文本，保留段落结构"""
        import fitz  # pymupdf
        texts = []
        with fitz.open(path) as doc:
            for page in doc:
                # get_text("text") 保留换行结构
                page_text = page.get_text("text")
                if page_text.strip():
                    texts.append(page_text)
        return "\n\n".join(texts)

    def _parse_docx(self, path: str) -> str:
        """python-docx 解析 DOCX，逐段提取"""
        from docx import Document as DocxDocument
        doc   = DocxDocument(path)
        parts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)
        # 也提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts)

    def _parse_html(self, path: str) -> str:
        """标准库 html.parser 提取正文，跳过 script/style"""
        from html.parser import HTMLParser

        class _Extractor(HTMLParser):
            def __init__(self):
                super().__init__()
                self.parts: List[str] = []
                self._skip = False

            def handle_starttag(self, tag, attrs):
                if tag in ("script", "style", "head", "nav", "footer"):
                    self._skip = True

            def handle_endtag(self, tag):
                if tag in ("script", "style", "head", "nav", "footer"):
                    self._skip = False

            def handle_data(self, data):
                if not self._skip and data.strip():
                    self.parts.append(data.strip())

        p = _Extractor()
        p.feed(Path(path).read_text("utf-8", errors="ignore"))
        return "\n".join(p.parts)

    def _parse_text(self, path: str) -> str:
        """纯文本直接读取"""
        return Path(path).read_text("utf-8", errors="ignore")


# ── 文本清洗 ──────────────────────────────────

class TextCleaner:
    """清洗：去除多余空白、控制字符，保留中英文内容"""

    # 保留：ASCII可打印 + 中文 + 中文标点 + 全角符号
    _KEEP_PATTERN = re.compile(
        r"[^\x09\x0a\x0d\x20-\x7e"
        r"\u4e00-\u9fa5"       # 中文
        r"\u3000-\u303f"       # 中文标点
        r"\uff00-\uffef"       # 全角字符
        r"\u2014-\u2027]"      # 破折号等
    )

    def clean(self, text: str) -> str:
        if not text:
            return ""
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = re.sub(r"[ \t]{3,}", " ", text)
        text = self._KEEP_PATTERN.sub("", text)
        return text.strip()


# ── 滑动窗口分块 ──────────────────────────────

class TextSplitter:
    """在句末断开的滑动窗口分块"""

    # 优先断句符（中英文）
    SENTENCE_ENDS = ["。", "！", "？", "\n\n", ".", "!", "?", "\n", "；", ";"]

    def __init__(self, chunk_size: int = None, overlap: int = None):
        self.chunk_size = chunk_size or settings.CHUNK_SIZE
        self.overlap    = max(0, overlap if overlap is not None else settings.CHUNK_OVERLAP)

    def split(self, text: str) -> List[str]:
        if not text or not text.strip():
            return []
        # 文本比 chunk_size 短，直接返回一块
        if len(text) <= self.chunk_size:
            stripped = text.strip()
            return [stripped] if stripped else []
        chunks = []
        start  = 0
        n      = len(text)
        while start < n:
            end = min(start + self.chunk_size, n)
            # 在中间点之后寻找句末断点
            if end < n:
                mid = start + self.chunk_size // 2
                for sep in self.SENTENCE_ENDS:
                    pos = text.rfind(sep, mid, end)
                    if pos != -1:
                        end = pos + len(sep)
                        break
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            # 下一块起点（带重叠）
            next_start = end - self.overlap
            if next_start <= start:
                next_start = start + 1
            start = next_start
        return chunks


# ── 质量评估 ──────────────────────────────────

class QualityChecker:
    """基于有效 chunk 比例和平均长度评估文档质量"""

    MIN_VALID_LEN = 10  # 最短有效 chunk 字符数（中文10字即有效内容）

    def evaluate(self, chunks: List[str]) -> Dict[str, Any]:
        if not chunks:
            return {"score": 0.0, "valid_ratio": 0.0, "avg_length": 0.0, "total": 0, "valid": 0}
        valid   = [c for c in chunks if len(c.strip()) >= self.MIN_VALID_LEN]
        avg_len = sum(len(c) for c in chunks) / len(chunks)
        len_score = min(avg_len / 100.0, 1.0)   # 平均长度分（100字为满分，适合中文短句）
        ratio_score = len(valid) / len(chunks)
        score = ratio_score * 0.7 + len_score * 0.3
        return {
            "score":       round(score, 4),
            "valid_ratio": round(ratio_score, 4),
            "avg_length":  round(avg_len, 1),
            "total":       len(chunks),
            "valid":       len(valid),
        }


# ── 文档处理服务主入口 ────────────────────────

class DocumentService:

    def __init__(self):
        self.parser   = DocParser()
        self.cleaner  = TextCleaner()
        self.splitter = TextSplitter()
        self.checker  = QualityChecker()

    async def process(self, doc_id: str, file_path: str, db: AsyncSession):
        logger.info(f"开始处理文档 doc_id={doc_id} path={file_path}")
        await self._set_status(db, doc_id, "processing")

        try:
            # Step 1: 解析
            raw = self.parser.parse(file_path)
            if not raw or not raw.strip():
                raise ValueError("文档内容为空，无法解析")

            # Step 2: 清洗
            text = self.cleaner.clean(raw)
            if not text:
                raise ValueError("清洗后内容为空")

            # Step 3: 分块
            chunks = self.splitter.split(text)
            if not chunks:
                raise ValueError("分块结果为空")

            # Step 4: 质量评估
            quality = self.checker.evaluate(chunks)
            logger.info(
                f"文档 {doc_id} 质量评估: score={quality['score']:.3f} "
                f"valid={quality['valid']}/{quality['total']} avg_len={quality['avg_length']:.0f}"
            )
            if quality["score"] < settings.QUALITY_THRESHOLD:
                raise ValueError(
                    f"质量分 {quality['score']:.2f} 低于阈值 {settings.QUALITY_THRESHOLD}，拒绝入库"
                )

            # Step 5: Embedding（批量，带进度日志）
            logger.info(f"开始 Embedding {len(chunks)} 个 chunk...")
            embeddings = await embed_client.embed_batch(chunks)

            # Step 6: Milvus 向量入库
            chunk_ids = [str(uuid.uuid4()) for _ in chunks]
            milvus_db.insert(
                ids        = chunk_ids,
                doc_ids    = [doc_id] * len(chunks),
                chunk_idxs = list(range(len(chunks))),
                texts      = chunks,
                embeddings = embeddings,
            )

            # Step 7: 更新 BM25 内存索引
            from app.rag.retriever import retriever
            retriever.add_texts(chunks)

            # Step 8: PostgreSQL 持久化 Chunk 记录
            db.add_all([
                Chunk(
                    id         = chunk_ids[i],
                    doc_id     = doc_id,
                    content    = chunks[i],
                    chunk_idx  = i,
                    char_count = len(chunks[i]),
                    meta_info  = {"doc_id": doc_id, "idx": i},
                )
                for i in range(len(chunks))
            ])

            # Step 9: 更新文档状态为 done
            await db.execute(
                update(Document).where(Document.id == doc_id).values(
                    status      = "done",
                    parse_score = quality["score"],
                    chunk_count = len(chunks),
                )
            )
            await db.commit()
            logger.info(f"✅ 文档 {doc_id} 处理完成，共 {len(chunks)} 个 chunk")

        except Exception as e:
            logger.error(f"❌ 文档 {doc_id} 处理失败: {e}")
            try:
                await db.rollback()
                await self._set_status(db, doc_id, "failed", str(e))
            except Exception as inner:
                logger.error(f"状态回滚失败: {inner}")
            raise

    async def _set_status(
        self, db: AsyncSession, doc_id: str, status: str, error: str = None
    ):
        vals: Dict[str, Any] = {"status": status}
        if error:
            vals["error_msg"] = error[:500]
        try:
            await db.execute(
                update(Document).where(Document.id == doc_id).values(**vals)
            )
            await db.commit()
        except Exception as e:
            logger.error(f"_set_status 失败: {e}")


doc_service = DocumentService()
